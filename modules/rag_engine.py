"""RAG Engine for O2C AI Monitor"""
import os
import sys
import json
import pickle
import hashlib
import numpy as np
import pandas as pd
from pathlib import Path
from datetime import datetime
from typing import List, Dict, Optional

# UTF-8 stdout encoding for Windows console
if sys.platform == "win32":
    try:
        sys.stdout.reconfigure(encoding="utf-8", errors="replace")
        sys.stderr.reconfigure(encoding="utf-8", errors="replace")
    except AttributeError:
        pass

# Runtime check for RAG packages - try to import them now
try:
    from sentence_transformers import SentenceTransformer
    import faiss
    from pypdf import PdfReader
    import docx
    import openpyxl
    RAG_AVAILABLE = True
except ImportError:
    RAG_AVAILABLE = False
    # Define dummy classes so module can still load
    SentenceTransformer = None
    faiss = None

def _check_rag_available():
    """Re-check RAG availability at runtime (in case packages were installed after import)"""
    try:
        # Try importing fresh to bypass any cached state
        import importlib
        import sys
        
        # Force fresh import
        if 'sentence_transformers' in sys.modules:
            importlib.reload(sys.modules['sentence_transformers'])
        if 'faiss' in sys.modules:
            importlib.reload(sys.modules['faiss'])
            
        from sentence_transformers import SentenceTransformer
        import faiss
        from pypdf import PdfReader
        import docx
        import openpyxl
        
        # Update globals if successful
        globals()['SentenceTransformer'] = SentenceTransformer
        globals()['faiss'] = faiss
        globals()['PdfReader'] = PdfReader
        globals()['docx'] = docx
        globals()['openpyxl'] = openpyxl
        globals()['RAG_AVAILABLE'] = True
        
        return True
    except ImportError:
        return False

# Use relative import for package-based execution
try:
    from .config import DOCS_DIR, VECTOR_DIR, CHUNKS_DIR, EMBEDDING_MODEL, CHUNK_SIZE, CHUNK_OVERLAP, TOP_K_RESULTS
except ImportError:
    # Fallback for direct execution from modules directory
    from config import DOCS_DIR, VECTOR_DIR, CHUNKS_DIR, EMBEDDING_MODEL, CHUNK_SIZE, CHUNK_OVERLAP, TOP_K_RESULTS


class DocumentLoader:
    """Loads documents from multiple formats"""
    
    SUPPORTED = [".txt", ".pdf", ".docx", ".csv", ".xlsx"]

    def __init__(self, docs_dir: Path = DOCS_DIR):
        self.docs_dir = docs_dir

    def load_all(self) -> List[Dict]:
        """Load all supported documents recursively from all subdirectories"""
        # Recursively find all files
        all_files = []
        for root, dirs, files in os.walk(self.docs_dir):
            for file in files:
                file_path = Path(root) / file
                if file_path.suffix.lower() in self.SUPPORTED:
                    all_files.append(file_path)
        
        print(f"📂 Loading {len(all_files)} documents (scanning all folders)...")
        docs = []
        for f in all_files:
            # Get relative path from docs_dir for better display
            rel_path = f.relative_to(self.docs_dir)
            print(f"   📄 {rel_path}...", end=" ")
            
            text = self._load_one(f)
            if text:
                # Use folder name as category
                folder_name = f.parent.name if f.parent != self.docs_dir else "general"
                category = self._normalize_category(folder_name)
                
                docs.append({
                    "filename": f.name,
                    "folder": folder_name,
                    "category": category,
                    "relative_path": str(rel_path),
                    "text": text,
                    "char_count": len(text)
                })
                print(f"✅ {len(text)} chars")
            else:
                print("❌ Failed")
        
        print(f"\n✅ Loaded {len(docs)} documents from all folders\n")
        return docs

    def _load_one(self, file_path: Path) -> Optional[str]:
        try:
            ext = file_path.suffix.lower()
            if ext == ".txt":
                return file_path.read_text(encoding="utf-8", errors="ignore")
            elif ext == ".pdf":
                reader = PdfReader(file_path)
                return "\n".join(page.extract_text() or "" for page in reader.pages)
            elif ext == ".docx":
                try:
                    doc = docx.Document(file_path)
                    p_texts = [p.text for p in doc.paragraphs if p.text]
                    t_texts = []
                    for table in doc.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                if cell.text.strip():
                                    t_texts.append(cell.text.strip())
                    full_text = "\n".join(p_texts + t_texts)
                    if full_text.strip():
                        return full_text
                except Exception:
                    pass
                
                # Direct XML extraction fallback
                try:
                    import zipfile
                    import xml.etree.ElementTree as ET
                    with zipfile.ZipFile(file_path) as z:
                        xml_content = z.read('word/document.xml')
                        root = ET.fromstring(xml_content)
                        texts = [node.text for node in root.iter('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t') if node.text]
                        if texts:
                            return "\n".join(texts)
                except Exception:
                    pass
                return None
            elif ext == ".csv":
                df = pd.read_csv(file_path)
                return df.to_string()
            elif ext == ".xlsx":
                df = pd.read_excel(file_path)
                return df.to_string()
        except Exception as e:
            print(f"Error loading {file_path.name}: {e}")
            return None

    def _normalize_category(self, folder_name: str) -> str:
        """Convert folder name to a normalized category slug"""
        # Convert "Clinic SLA's" -> "clinic_slas"
        # Convert "History Resolution Logs" -> "history_resolution_logs"
        normalized = folder_name.lower()
        # Remove possessives
        normalized = normalized.replace("'s", "s").replace("'", "")
        # Replace spaces with underscores
        normalized = normalized.replace(" ", "_")
        # Remove special characters
        normalized = ''.join(c for c in normalized if c.isalnum() or c == '_')
        return normalized


class ClauseAwareChunker:
    """Splits documents preserving structural sections, contract clauses, and headers"""

    def __init__(self, chunk_size: int = CHUNK_SIZE, chunk_overlap: int = CHUNK_OVERLAP):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap

    def chunk_documents(self, docs: List[Dict]) -> List[Dict]:
        """Chunk all documents using clause-aware boundary detection"""
        print("✂️  Chunking documents (Clause & Section-Aware)...")
        all_chunks = []
        for doc in docs:
            chunks = self._chunk_one(
                text=doc["text"],
                filename=doc["filename"],
                category=doc["category"],
                folder=doc.get("folder", "general"),
                relative_path=doc.get("relative_path", doc["filename"])
            )
            all_chunks.extend(chunks)
            rel_path = doc.get("relative_path", doc["filename"])
            print(f"   📄 {rel_path:<50} → {len(chunks)} chunks")
        
        print(f"\n✅ Total chunks created: {len(all_chunks)}")
        return all_chunks
    
    def save_chunks(self, chunks: List[Dict]):
        """Save chunks to disk for inspection"""
        chunks_file = CHUNKS_DIR / "all_chunks.json"
        with open(chunks_file, "w", encoding="utf-8") as f:
            json.dump(chunks, f, indent=2, ensure_ascii=False)
        print(f"💾 Chunks saved to: {chunks_file}")

    def _chunk_one(self, text: str, filename: str, category: str, folder: str = "general", relative_path: str = None) -> List[Dict]:
        import re
        chunks = []
        chunk_id = 0
        
        # Split on standard clause/heading patterns (\n\n, 1. , 2. , 5.1 , TICKET, SECTION)
        raw_sections = re.split(r'(?=\n(?:[0-9]+\.[0-9]*\s+|TICKET\s+|SECTION\s+|INC-[0-9]+\s+|[A-Z\s]{4,}:))|\n\n+', text)
        
        current_chunk = ""
        header_context = filename.replace(".docx", "").replace(".txt", "")
        
        for section in raw_sections:
            sec_clean = section.strip()
            if not sec_clean:
                continue
            
            # If section fits in chunk size, aggregate
            if len(current_chunk) + len(sec_clean) <= self.chunk_size:
                current_chunk = (current_chunk + "\n\n" + sec_clean).strip()
            else:
                if current_chunk:
                    chunks.append({
                        "chunk_id": f"{filename}_{chunk_id}",
                        "filename": filename,
                        "folder": folder,
                        "category": category,
                        "relative_path": relative_path or filename,
                        "text": current_chunk,
                        "doc_title": header_context
                    })
                    chunk_id += 1
                
                # If individual section is longer than chunk_size, split by sentences
                if len(sec_clean) > self.chunk_size:
                    sentences = re.split(r'(?<=[.!?])\s+', sec_clean)
                    sub_chunk = ""
                    for s in sentences:
                        if len(sub_chunk) + len(s) <= self.chunk_size:
                            sub_chunk = (sub_chunk + " " + s).strip()
                        else:
                            if sub_chunk:
                                chunks.append({
                                    "chunk_id": f"{filename}_{chunk_id}",
                                    "filename": filename,
                                    "folder": folder,
                                    "category": category,
                                    "relative_path": relative_path or filename,
                                    "text": sub_chunk,
                                    "doc_title": header_context
                                })
                                chunk_id += 1
                            sub_chunk = s
                    current_chunk = sub_chunk
                else:
                    current_chunk = sec_clean
        
        if current_chunk.strip():
            chunks.append({
                "chunk_id": f"{filename}_{chunk_id}",
                "filename": filename,
                "folder": folder,
                "category": category,
                "relative_path": relative_path or filename,
                "text": current_chunk.strip(),
                "doc_title": header_context
            })
        
        return chunks


# Backwards compatibility alias
TextChunker = ClauseAwareChunker


class BM25Index:
    """Self-contained Okapi BM25 Sparse Keyword Index"""

    def __init__(self, k1: float = 1.5, b: float = 0.75):
        self.k1 = k1
        self.b = b
        self.corpus = []
        self.doc_len = []
        self.avg_doc_len = 0.0
        self.doc_freqs = []
        self.idf = {}
        self.metadata = []

    def _tokenize(self, text: str) -> List[str]:
        import re
        tokens = re.findall(r'[a-zA-Z0-9$_\-%]+', text.lower())
        return [t for t in tokens if len(t) > 1]

    def build_index(self, chunks: List[Dict]):
        """Build BM25 index from text chunks"""
        import math
        self.metadata = chunks
        self.corpus = [self._tokenize(c["text"]) for c in chunks]
        self.doc_len = [len(doc) for doc in self.corpus]
        total_tokens = sum(self.doc_len)
        self.avg_doc_len = (total_tokens / len(self.corpus)) if self.corpus else 1.0

        # Term document frequency
        df = {}
        for doc in self.corpus:
            unique_terms = set(doc)
            for term in unique_terms:
                df[term] = df.get(term, 0) + 1

        # Calculate standard IDF
        n_docs = len(self.corpus)
        self.idf = {}
        for term, freq in df.items():
            self.idf[term] = math.log(1.0 + (n_docs - freq + 0.5) / (freq + 0.5))

    def search(self, query: str, top_k: int = TOP_K_RESULTS, category: str = None) -> List[Dict]:
        """Search BM25 index for sparse keyword matches"""
        q_tokens = self._tokenize(query)
        if not q_tokens or not self.corpus:
            return []

        scores = []
        for idx, doc in enumerate(self.corpus):
            chunk = self.metadata[idx]
            if category is not None and chunk.get("category") != category:
                continue

            doc_l = self.doc_len[idx]
            score = 0.0
            term_counts = {}
            for t in doc:
                term_counts[t] = term_counts.get(t, 0) + 1

            for q in q_tokens:
                if q in term_counts:
                    tf = term_counts[q]
                    idf = self.idf.get(q, 0.0)
                    denom = tf + self.k1 * (1.0 - self.b + self.b * (doc_l / self.avg_doc_len))
                    score += idf * (tf * (self.k1 + 1.0)) / denom

            if score > 0.0:
                scores.append((idx, score))

        scores.sort(key=lambda x: x[1], reverse=True)
        results = []
        max_score = scores[0][1] if scores else 1.0
        for rank, (idx, sc) in enumerate(scores[:top_k], 1):
            norm_sc = min(1.0, sc / max(1.0, max_score))
            results.append({
                **self.metadata[idx],
                "bm25_score": float(sc),
                "bm25_similarity": float(norm_sc),
                "bm25_rank": rank
            })
        return results


class VectorStore:
    """Embedding engine + FAISS vector store + BM25 Hybrid Fusion"""

    INDEX_FILE = VECTOR_DIR / "index.faiss"
    META_FILE = VECTOR_DIR / "metadata.pkl"
    CHUNKS_FILE = VECTOR_DIR / "chunks.pkl"
    BM25_FILE = VECTOR_DIR / "bm25.pkl"

    def __init__(self, model_name: str = EMBEDDING_MODEL):
        if not _check_rag_available():
            raise ImportError("RAG packages not installed. Install: sentence-transformers, faiss-cpu, pypdf, python-docx, openpyxl")
        
        print(f"🤖 Loading embedding model: {model_name}")
        self.model = SentenceTransformer(model_name)
        try:
            self.dim = self.model.get_embedding_dimension()
        except AttributeError:
            self.dim = self.model.get_sentence_embedding_dimension()
        self.index = None
        self.metadata = []
        self.bm25 = BM25Index()
        print(f"✅ Model loaded | Embedding dimension: {self.dim}")

    def build_index(self, chunks: List[Dict]):
        """Build FAISS dense vector index and BM25 sparse keyword index"""
        print(f"\n🔨 Building Hybrid vector & BM25 index for {len(chunks)} chunks...")
        
        texts = [c["text"] for c in chunks]
        print("   Generating normalized dense embeddings...")
        embeddings = self.model.encode(texts, show_progress_bar=True, normalize_embeddings=True)
        
        print("   Building FAISS Cosine Similarity index...")
        dim = embeddings.shape[1]
        self.index = faiss.IndexFlatIP(dim)
        self.index.add(embeddings.astype("float32"))
        self.metadata = chunks
        
        print("   Building Okapi BM25 sparse index...")
        self.bm25.build_index(chunks)
        
        print(f"\n💾 Saving index files...")
        faiss.write_index(self.index, str(self.INDEX_FILE))
        with open(self.META_FILE, "wb") as f:
            pickle.dump(self.metadata, f)
        with open(self.CHUNKS_FILE, "wb") as f:
            pickle.dump(chunks, f)
        with open(self.BM25_FILE, "wb") as f:
            pickle.dump(self.bm25, f)
        
        print(f"✅ Hybrid Index saved: {len(chunks)} vectors + BM25 lexicon")

    def load_index(self) -> bool:
        """Load existing hybrid index"""
        if not self.INDEX_FILE.exists() or not self.META_FILE.exists():
            return False
        
        print("💾 Loading existing Hybrid RAG index...")
        self.index = faiss.read_index(str(self.INDEX_FILE))
        with open(self.META_FILE, "rb") as f:
            self.metadata = pickle.load(f)
        if self.BM25_FILE.exists():
            with open(self.BM25_FILE, "rb") as f:
                self.bm25 = pickle.load(f)
        else:
            self.bm25.build_index(self.metadata)
        print(f"✅ Loaded {len(self.metadata)} vectors & BM25 index")
        return True

    def search_hybrid(self, query: str, top_k: int = TOP_K_RESULTS, category: str = None) -> List[Dict]:
        """
        Execute Hybrid Search combining Dense Vector Cosine Similarity
        and Sparse BM25 Keyword Matching via Reciprocal Rank Fusion (RRF).
        """
        if self.index is None:
            return []
        
        # 1. Dense Vector Search
        query_vec = self.model.encode([query], normalize_embeddings=True)
        distances, indices = self.index.search(query_vec.astype("float32"), top_k * 3)
        
        dense_results = {}
        for rank, (dist, idx) in enumerate(zip(distances[0], indices[0]), 1):
            if idx < len(self.metadata):
                chunk = self.metadata[idx]
                if category is None or chunk.get("category") == category:
                    dense_results[idx] = {
                        "rank": rank,
                        "cos_sim": max(0.0, min(1.0, float(dist)))
                    }

        # 2. Sparse BM25 Search
        bm25_matches = self.bm25.search(query, top_k=top_k * 3, category=category)
        bm25_results = {}
        for b in bm25_matches:
            # find index in metadata
            idx = next((i for i, c in enumerate(self.metadata) if c["chunk_id"] == b["chunk_id"]), None)
            if idx is not None:
                bm25_results[idx] = {
                    "rank": b["bm25_rank"],
                    "bm25_score": b["bm25_similarity"]
                }

        # 3. Reciprocal Rank Fusion (RRF)
        all_candidates = set(dense_results.keys()).union(set(bm25_results.keys()))
        rrf_scored = []
        k_rrf = 60.0

        for idx in all_candidates:
            d_rank = dense_results.get(idx, {}).get("rank", 999)
            b_rank = bm25_results.get(idx, {}).get("rank", 999)
            d_sim = dense_results.get(idx, {}).get("cos_sim", 0.0)
            b_sim = bm25_results.get(idx, {}).get("bm25_score", 0.0)

            # RRF formula
            rrf_score = (1.0 / (k_rrf + d_rank)) + (1.0 / (k_rrf + b_rank))
            # Blended semantic confidence
            blended_conf = max(d_sim, 0.70 * d_sim + 0.30 * b_sim)

            chunk = self.metadata[idx]
            rrf_scored.append({
                **chunk,
                "score": float(blended_conf),
                "similarity": float(blended_conf),
                "rrf_score": float(rrf_score),
                "dense_cosine": float(d_sim),
                "bm25_score": float(b_sim)
            })

        # Sort by RRF rank score
        rrf_scored.sort(key=lambda x: x["rrf_score"], reverse=True)
        return rrf_scored[:top_k]

    def search(self, query: str, top_k: int = TOP_K_RESULTS, category: str = None) -> List[Dict]:
        """Search interface - uses Hybrid Search by default"""
        return self.search_hybrid(query=query, top_k=top_k, category=category)


class RAGQueryEngine:
    """Main RAG interface"""

    def __init__(self, vector_store: VectorStore):
        self.vs = vector_store

    def query(self, question: str, category: str = None, top_k: int = TOP_K_RESULTS, verbose: bool = True) -> Dict:
        """Main query method - matches notebook interface"""
        result = self.ask(question, category, top_k)
        if verbose:
            print(f"\n❔ Question: {question}")
            print(f"\n💡 Answer:\n{result['answer']}\n")
            print(f"📚 Sources ({len(result['sources'])})")
            for i, src in enumerate(result['sources'][:3], 1):
                print(f"   {i}. {src['filename']} ({src['category']}) | similarity: {src['similarity']:.2f}")
        return result
    
    def ask(self, question: str, category: str = None, top_k: int = TOP_K_RESULTS) -> Dict:
        """Ask a question and get answer with sources"""
        # Search for relevant chunks
        chunks = self.vs.search(question, top_k=top_k, category=category)
        
        if not chunks:
            return {
                "question": question,
                "answer": "No relevant information found.",
                "sources": [],
                "confidence": 0.0
            }
        
        # Build context from chunks
        context = "\n\n".join([f"[{c['filename']}]\n{c['text']}" for c in chunks])
        
        # Simple answer generation (can be enhanced with LLM)
        answer = self._generate_answer(question, context, chunks)
        
        # Extract sources
        sources = [
            {"filename": c["filename"], "category": c["category"], 
             "similarity": c["similarity"], "chunk_id": c["chunk_id"]}
            for c in chunks
        ]
        
        return {
            "question": question,
            "answer": answer,
            "sources": sources,
            "confidence": chunks[0]["similarity"] if chunks else 0.0,
            "context": context
        }

    def _generate_answer(self, question: str, context: str, chunks: List[Dict]) -> str:
        """Synthesize structured answer across top matching chunks"""
        if not chunks:
            return "No information found."
        
        # Aggregate unique key excerpts from top chunks
        excerpts = []
        seen_texts = set()
        for i, c in enumerate(chunks[:3], 1):
            text_snippet = c['text'].strip()
            # Clean up duplicates
            clean_snippet = " ".join(text_snippet.split())
            if clean_snippet not in seen_texts:
                seen_texts.add(clean_snippet)
                rel = c.get('relative_path', c['filename'])
                cat = c.get('category', 'general')
                sim = c.get('similarity', 0.0)
                excerpts.append(f"📄 [{rel}] (Category: {cat} | Similarity: {sim:.2f}):\n{text_snippet}")
        
        body = "\n\n".join(excerpts)
        
        sources_list = "\n".join([
            f"- {c.get('relative_path', c['filename'])} ({c.get('category', 'general')}) | Similarity: {c.get('similarity', 0.0):.2f}"
            for c in chunks[:4]
        ])
        
        return f"""{body}

📚 Referenced Policy Sources:
{sources_list}"""


class RAGEngine:
    """Complete RAG system - main interface"""

    def __init__(self):
        # Re-check availability at runtime (not cached import-time check)
        if not _check_rag_available():
            raise ImportError("RAG packages not available. Install: sentence-transformers, faiss-cpu, etc.")
        
        self.loader = DocumentLoader()
        self.chunker = ClauseAwareChunker()
        self.vector_store = VectorStore()
        self.query_engine = None

    def initialize(self, force_rebuild: bool = False) -> bool:
        """Initialize RAG system"""
        # Try to load existing index
        if not force_rebuild and self.vector_store.load_index():
            self.query_engine = RAGQueryEngine(self.vector_store)
            print("✅ RAG Engine initialized (loaded existing index)")
            return True
        
        # Build new index
        print("🛠️ Building new RAG index...")
        docs = self.loader.load_all()
        if not docs:
            print("⚠️  No documents found. Add documents to:", DOCS_DIR)
            return False
        
        chunks = self.chunker.chunk_documents(docs)
        self.chunker.save_chunks(chunks)  # Save chunks for inspection
        self.vector_store.build_index(chunks)
        self.query_engine = RAGQueryEngine(self.vector_store)
        print("✅ RAG Engine initialized (built new index)")
        return True

    def query(self, question: str, category: str = None, verbose: bool = True) -> Dict:
        """Query the RAG system (verbose output)"""
        if self.query_engine is None:
            return {"error": "RAG not initialized. Call initialize() first."}
        return self.query_engine.query(question, category=category, verbose=verbose)
    
    def ask(self, question: str, category: str = None) -> Dict:
        """Ask a question (returns dict only, no printing)"""
        if self.query_engine is None:
            return {"error": "RAG not initialized. Call initialize() first."}
        return self.query_engine.ask(question, category=category)