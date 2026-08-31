"""Strike Intelligence Generator - Converts strike news into RAG-ready policy documents"""
import sqlite3
import json
from pathlib import Path
from datetime import datetime
from typing import List, Dict
from collections import Counter

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    from .config import DB_PATH, DOCS_DIR
except ImportError:
    from config import DB_PATH, DOCS_DIR

try:
    from .ollama_service import OllamaService
except ImportError:
    from ollama_service import OllamaService


class StrikeIntelligenceGenerator:
    """
    Converts strike/disruption news into intelligence briefs for RAG.
    
    Purpose:
    - Strike news is NOT for validating RAG
    - Strike intelligence is RAG knowledge for route planning
    - Used by Copilot to avoid disrupted zones, find alternatives
    """
    
    def __init__(self, db_path=str(DB_PATH), output_dir=None):
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx not installed. Run: pip install python-docx")
        
        self.db_path = db_path
        self.output_dir = output_dir or (DOCS_DIR / "Strike_Intelligence")
        self.output_dir.mkdir(parents=True, exist_ok=True)
        self.ollama = OllamaService()
        if self.ollama.is_available():
            print(f"🤖 Connected to Local Ollama LLM ({self.ollama.model}) for AI Intelligence Synthesis")
    
    def generate_all_intelligence(self) -> List[str]:
        """Generate strike intelligence documents from all news articles"""
        print("\n" + "="*80)
        print("🚨 STRIKE INTELLIGENCE GENERATOR")
        print("="*80)
        
        articles = self._fetch_strike_articles()
        print(f"\n📊 Found {len(articles)} strike articles in database")
        
        if not articles:
            print("⚠️  No strike articles found. Run the main pipeline first.")
            return []
        
        intel_by_city = self._group_articles_by_city(articles)
        intel_by_category = self._group_articles_by_category(articles)
        
        generated_files = []
        print(f"\n📝 Generating intelligence briefs...")
        
        for city, city_articles in intel_by_city.items():
            if len(city_articles) >= 3:
                doc_path = self._create_city_strike_brief(city, city_articles)
                generated_files.append(doc_path)
                print(f"   ✅ {doc_path.name}")
        
        for category, cat_articles in intel_by_category.items():
            if len(cat_articles) >= 5:
                doc_path = self._create_category_strike_brief(category, cat_articles)
                generated_files.append(doc_path)
                print(f"   ✅ {doc_path.name}")
        
        master_doc = self._create_master_disruption_intelligence(articles)
        generated_files.append(master_doc)
        print(f"   ✅ {master_doc.name}")
        
        print(f"\n✅ Generated {len(generated_files)} intelligence documents")
        print(f"📁 Saved to: {self.output_dir}")
        print("="*80 + "\n")
        
        return [str(f) for f in generated_files]
    
    def _fetch_strike_articles(self) -> List[Dict]:
        conn = sqlite3.connect(self.db_path)
        conn.row_factory = sqlite3.Row
        cursor = conn.cursor()
        
        query = """SELECT title, city_mentioned as matched_cities, strike_type as category, 
                   published_date as published, source_name as source
                   FROM strike_news ORDER BY published_date DESC"""
        
        cursor.execute(query)
        articles = [dict(row) for row in cursor.fetchall()]
        conn.close()
        
        for article in articles:
            # city_mentioned is a single string, convert to list
            city = article.get('matched_cities')
            if city and city != 'Unknown':
                article['matched_cities'] = [city]
            else:
                article['matched_cities'] = []
        return articles
    
    def _group_articles_by_city(self, articles: List[Dict]) -> Dict[str, List[Dict]]:
        by_city = {}
        for article in articles:
            for city in article.get('matched_cities', []):
                if city not in by_city:
                    by_city[city] = []
                by_city[city].append(article)
        return by_city
    
    def _group_articles_by_category(self, articles: List[Dict]) -> Dict[str, List[Dict]]:
        by_category = {}
        for article in articles:
            category = article.get('category', 'general')
            if category not in by_category:
                by_category[category] = []
            by_category[category].append(article)
        return by_category
    
    def _create_city_strike_brief(self, city: str, articles: List[Dict]) -> Path:
        doc = Document()
        
        # Telemetry & Incident Aggregation
        high_sev = [a for a in articles if 'HIGH' in str(a.get('severity', '')).upper()]
        med_sev = [a for a in articles if 'MEDIUM' in str(a.get('severity', '')).upper()]
        low_sev = [a for a in articles if 'LOW' in str(a.get('severity', '')).upper() or not a.get('severity')]
        city_code = city[:3].upper()
        
        # Title & Metadata Block
        title = doc.add_heading(f'{city.upper()} TRANSPORTATION DISRUPTION INTELLIGENCE & ROUTING PLAYBOOK', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph(
            f"Monitored Freight Hub: {city} Logistics Cluster\n"
            f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')} | Total Disruption Incidents Analyzed: {len(articles)}\n"
            f"Disruption Severity Breakdown: 🔴 HIGH: {len(high_sev)} | 🟡 MEDIUM: {len(med_sev)} | 🟢 LOW: {len(low_sev)}"
        )
        
        # 1. Extracted Disruption Incident Registry (Table)
        doc.add_heading('1. EXTRACTED DISRUPTION INCIDENT REGISTRY', 1)
        
        table = doc.add_table(rows=1, cols=5)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Incident ID & Date'
        hdr_cells[1].text = 'Source'
        hdr_cells[2].text = 'Disruption Modality'
        hdr_cells[3].text = 'Stated Trigger / Demand'
        hdr_cells[4].text = 'Freight Impact Severity'
        
        for idx, art in enumerate(articles[:10], 1):
            row_cells = table.add_row().cells
            row_cells[0].text = f"INC-{city_code}-{idx:02d}\n({art.get('published', 'N/A')})"
            row_cells[1].text = art.get('source', 'News Wire')
            row_cells[2].text = art.get('category', 'General').title()
            
            # Extract clean trigger
            clean_title = art.get('title', 'Transport Strike')
            if ' - ' in clean_title:
                clean_title = clean_title.split(' - ')[0]
            row_cells[3].text = clean_title[:90]
            
            sev = art.get('severity', 'LOW')
            row_cells[4].text = f"🔴 HIGH" if 'HIGH' in str(sev).upper() else (f"🟡 MEDIUM" if 'MEDIUM' in str(sev).upper() else f"🟢 LOW")
            
        doc.add_paragraph()  # Spacing
        
        # 2. Critical Bottlenecks & Highway Bypass Directives
        doc.add_heading('2. CRITICAL BOTTLENECKS & HIGHWAY BYPASS DIRECTIVES', 1)
        
        corridor_info = {
            "Mumbai": (
                "• Primary Choke Points: NH-48 (Mumbai-Pune Expressway), JNPT Port Container Freight Stations, Bhiwandi Central Warehouse Cluster.\n"
                "• Recommended FTL Bypass: Divert long-haul freight via Eastern Freeway & JNPT coastal corridor; utilize rail container drayage."
            ),
            "Delhi": (
                "• Primary Choke Points: NH-44 Kundli Border (Delhi-Haryana entry), Sanjay Gandhi Transport Nagar (SGTN), DND Flyway.\n"
                "• Recommended FTL Bypass: Utilize Kundli-Manesar-Palwal (KMP) Expressway to bypass inner Delhi urban commercial vehicle restrictions."
            ),
            "Bangalore": (
                "• Primary Choke Points: NH-44 (Hosur Road / Electronic City toll), Nelamangala Logistics Hub, Peenya Industrial Gate.\n"
                "• Recommended FTL Bypass: Route southern linehauls via NICE Road ring bypass; avoid inner Ring Road during peak hours (08:00-20:00)."
            ),
            "Chennai": (
                "• Primary Choke Points: NH-16 (Chennai-Kolkata corridor), Sriperumbudur Industrial Corridor, Chennai Port Trust Gates.\n"
                "• Recommended FTL Bypass: Divert northern freight via Chennai Outer Ring Road (CORR) and Minjur bypass."
            ),
            "Kolkata": (
                "• Primary Choke Points: NH-19 (Durgapur Expressway), Dankuni Freight Terminal, Vidyasagar Setu approach.\n"
                "• Recommended FTL Bypass: Route heavy commercial vehicles via Kona Expressway and Belghoria bypass."
            ),
            "Hyderabad": (
                "• Primary Choke Points: NH-65 (Hyderabad-Vijayawada), Outer Ring Road (ORR) exits 11-14, Shamshabad cargo terminal.\n"
                "• Recommended FTL Bypass: Utilize full ORR loop bypass; avoid inner arterial corridors during daytime transit curfews."
            ),
            "Pune": (
                "• Primary Choke Points: NH-48 (Pune-Bangalore Highway), Chakan Industrial Belt, Talegaon Logistics Park.\n"
                "• Recommended FTL Bypass: Route via Pune Outer Ring corridor and Talegaon-Chakan road off-peak."
            )
        }
        
        bypass_desc = corridor_info.get(city, (
            f"• Primary Choke Points: {city} primary national highway arterial entry gates and regional transshipment yards.\n"
            f"• Recommended FTL Bypass: Reroute freight via state highway ring corridors; enforce night dispatch schedules."
        ))
        doc.add_paragraph(bypass_desc)
        
        # 3. Autonomous Copilot Adjudication & Legal Rules (Discrete Rules)
        doc.add_heading('3. AUTONOMOUS COPILOT ADJUDICATION & LEGAL RULES', 1)
        
        rules = [
            f"[RULE-S-{city_code}-01] FORCE MAJEURE & SLA PENALTY WAIVER (SECTION 8.4):\n"
            f"  Government-declared bandhs or verified labor road blockades >12h in {city} grant 100% carrier SLA delay penalty waiver ($500/day -> $0.00) upon submission of GPS telematics or police traffic advisory.",
            
            f"[RULE-S-{city_code}-02] EMERGENCY AIR FREIGHT & MODE SHIFT (CLAUSE 2.3):\n"
            f"  For life-critical veterinary diets and urgent pharmaceuticals (sap_mara.specialty_diet_flag = 1) stranded >24h in {city}, authorize immediate $1,000 Air Freight replacement per Medical Stock-Out Policy 2024-04.",
            
            f"[RULE-S-{city_code}-03] TRUCK DETENTION & DEMURRAGE LIABILITY CAP:\n"
            f"  Carrier demurrage and truck detention claims during labor strikes are strictly capped at $100.00/day per vehicle, requiring verified GPS geofence timestamps.",
            
            f"[RULE-S-{city_code}-04] DYNAMIC TRANSIT BUFFER & CLINIC RESCHEDULING:\n"
            f"  Automatically inject a +12.0h to +24.0h dynamic safety buffer into predicted arrival time (PDD). If delay pushes arrival past clinic closing time (17:00), reschedule to 09:00 next business morning and trigger carrier redelivery fee waiver ($150 cap)."
        ]
        
        for r in rules:
            doc.add_paragraph(r)
            
        # 4. Copilot Deterministic Action Checklist
        doc.add_heading('4. COPILOT DETERMINISTIC ACTION CHECKLIST', 1)
        checklist = [
            f"[x] Step 1: Scan active SAP linehauls (sap_vttk) scheduled to pass through {city} disruption zone.",
            f"[x] Step 2: Calculate revised ETA with [RULE-S-{city_code}-04] dynamic buffer (+12h to +24h).",
            f"[x] Step 3: If delay >24h for specialty diets, trigger [RULE-S-{city_code}-02] $1,000 Air Freight replacement.",
            f"[x] Step 4: Dispatch automated Actionable Adaptive Card to Regional Logistics Director for financial risk >$500."
        ]
        for item in checklist:
            doc.add_paragraph(item)
            
        # Save Document
        safe_city = city.replace('/', '-').replace('\\', '-')
        filename = f"{safe_city}_Strike_Intelligence.docx"
        doc_path = self.output_dir / filename
        doc.save(str(doc_path))
        return doc_path
    
    def _create_category_strike_brief(self, category: str, articles: List[Dict]) -> Path:
        doc = Document()
        cat_name = category.replace('_', ' ').title()
        title = doc.add_heading(f'{cat_name} Transportation Disruption Pattern & Risk Analysis', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph(f"Report Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        doc.add_paragraph(f"Transit Modality Category: {cat_name} Logistics Operations")
        doc.add_paragraph(f"Total Sector Incidents Analyzed: {len(articles)}")
        
        # 1. Modality Profile
        doc.add_heading('1. Sector Disruption Profile & Vulnerability', 1)
        doc.add_paragraph(
            f"This operational brief outlines systemic risk patterns, union actions, and transit vulnerabilities "
            f"specific to {cat_name} transportation across India. Disruptions in this sector typically stem from "
            f"toll tariff revisions, fuel price volatility, regulatory compliance enforcement (e.g. e-way bill disputes), "
            f"and driver union strikes."
        )
        
        # 2. Geographic Distribution
        all_cities = []
        for article in articles:
            all_cities.extend(article.get('matched_cities', []))
        city_impacts = Counter(all_cities)
        
        doc.add_heading('2. Geographic Impact Distribution', 1)
        for city, count in city_impacts.most_common(10):
            doc.add_paragraph(f"• {city}: {count} recorded incident(s)")
            
        # 3. Selected Real-World Incident Case Studies
        doc.add_heading('3. Key Sector Incidents & Historical Evidence', 1)
        for idx, art in enumerate(articles[:10], 1):
            p = doc.add_paragraph()
            p.add_run(f"Case #{idx}: {art.get('title', 'Unknown')}\n").bold = True
            p.add_run(f"• Date: {art.get('published', 'N/A')} | Source: {art.get('source', 'Unknown')}\n")
            if art.get('description'):
                clean_desc = art.get('description', '').replace('<a', '').replace('</a>', '').replace('&nbsp;', ' ')
                p.add_run(f"• Context: {clean_desc[:250]}...\n")
        
        # 4. Carrier Contract Adjudication & Mitigation
        doc.add_heading('4. Carrier SLA Adjudication & Operational Directives', 1)
        doc.add_paragraph(
            f"Standard operating procedures when {cat_name} disruptions impact Order-to-Cash deliveries:"
        )
        
        modality_guidance = {
            "truck": [
                "FTL Freight Rerouting: If primary highway corridor is blocked by truck strikes, carriers must attempt approved secondary state highway bypasses within 4 hours.",
                "Demurrage & Detention Caps: Carrier detention claims during nationwide truck strikes are capped at $100/day and require GPS geofence verification.",
                "Short-Dated Inventory Protection: For products with <60 days shelf-life stuck in highway transit, trigger Quality QA hold if transit exceeds 48 hours."
            ],
            "railway": [
                "Rail Yard Demurrage Rules: Rail container demurrage incurred due to labor strikes at inland container depots (ICDs) is reimbursable up to $500 per container upon railway receipt submission.",
                "Intermodal Drayage Shift: Transfer stranded rail cargo to dedicated road linehauls within 24 hours of rail stoppage."
            ],
            "bandh": [
                "Total Force Majeure Exemption: Complete Bharat Bandh or State Bandhs automatically activate Force Majeure Section 8.2 across all transit modes.",
                "Warehouse Dispatch Lockdown: No outbound shipments may be tendered during active bandh curfew hours to prevent en-route cargo damage or looting."
            ],
            "bus": [
                "Commuter Spillover Effect: State transport bus strikes cause heavy arterial road congestion (+2 to +4 hours linehaul delay). Planners should adjust dispatch times."
            ]
        }
        
        guidelines = modality_guidance.get(category.lower(), [
            "Assess transit delay duration against contractual delivery grace period (24 hours).",
            "Verify whether disruption is recognized by local transport authority as Force Majeure.",
            "Proactively notify destination clinics of revised ETA window."
        ])
        
        for g in guidelines:
            doc.add_paragraph(f"• {g}")
            
        # Sanitize category name for filename
        safe_cat = cat_name.replace('/', '-').replace('\\', '-')
        filename = f"{safe_cat}_Pattern_Analysis.docx"
        doc_path = self.output_dir / filename
        doc.save(str(doc_path))
        return doc_path
    
    def _create_master_disruption_intelligence(self, all_articles: List[Dict]) -> Path:
        doc = Document()
        title = doc.add_heading('Master Transportation Disruption Intelligence & National Playbook', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph(f"Report Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        doc.add_paragraph(f"National Scope: India Logistics Network & Veterinary Supply Chain")
        doc.add_paragraph(f"Total National Disruption Records Analyzed: {len(all_articles)}")
        
        doc.add_heading('1. National Supply Chain Risk Overview', 1)
        doc.add_paragraph(
            "This master intelligence brief synthesizes nationwide transportation labor disruptions, strike trends, "
            "and highway bottlenecks across India. It serves as the primary ground-truth knowledge base for the "
            "O2C Autonomous Delivery Risk Copilot (Phases 2, 4, and 5)."
        )
        
        # High Severity Summary
        high_sev = [a for a in all_articles if 'HIGH' in str(a.get('severity', '')).upper()]
        doc.add_heading('2. Critical National Disruptions (High Severity)', 1)
        doc.add_paragraph(f"Identified {len(high_sev)} major nationwide or state-wide disruption events:")
        
        for idx, art in enumerate(high_sev[:10], 1):
            p = doc.add_paragraph()
            p.add_run(f"• [{art.get('published', 'N/A')}] {art.get('title', 'Unknown')}\n").bold = True
            p.add_run(f"  Source: {art.get('source', 'Unknown')} | Region: {art.get('matched_cities', ['National'])[0]}\n")
            
        doc.add_heading('3. Autonomous AI Orchestration Decision Matrix', 1)
        actions = [
            "Phase 1 Ingestion Integration: Automatically query live weather and strike news on daily schedule; populate SQLite database and update vector knowledge base.",
            "Phase 2 RAG Retrieval: When an order risk is evaluated, retrieve matching city and modality intelligence to adjudicate Force Majeure eligibility and transit delay buffers.",
            "Phase 3 ML Synergy: Combine strike severity alerts with historical transit velocity to adjust predicted delay hours and probability.",
            "Phase 4 Multi-Agent Adjudication: Route Supervisor verifies alternative corridors; Contract Adjudicator calculates carrier chargebacks and SLA penalties ($500/day vs 5%/day); Quality Mitigation agent enforces cold-chain holds.",
            "Phase 5 ERP & Teams Action: Update SAP ERP delivery block / date (VDATU); dispatch Actionable Adaptive Card to Regional Logistics Director for approvals exceeding $500."
        ]
        for action in actions:
            doc.add_paragraph(action, style='List Bullet')
        
        doc_path = self.output_dir / "Master_Disruption_Intelligence.docx"
        doc.save(str(doc_path))
        return doc_path


if __name__ == "__main__":
    generator = StrikeIntelligenceGenerator()
    files = generator.generate_all_intelligence()
    print(f"\n📄 Generated files:")
    for f in files:
        print(f"   - {f}")