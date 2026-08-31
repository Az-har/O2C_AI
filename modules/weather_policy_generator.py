"""Weather Policy Generator - Converts weather data into RAG-ready policy documents"""
import sqlite3
import json
from pathlib import Path
from datetime import datetime
from typing import List, Dict
try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    from .config import DB_PATH, DOCS_DIR
except ImportError:
    from config import DB_PATH, DOCS_DIR

try:
    from .ollama_service import OllamaService
except ImportError:
    from ollama_service import OllamaService


class WeatherPolicyGenerator:
    """
    Converts weather alerts from database into policy documents for RAG.
    
    Purpose:
    - Weather data is NOT for validating RAG
    - Weather policies are RAG knowledge sources
    - Used by Copilot to apply Force Majeure, route changes, etc.
    """
    
    def __init__(self, db_path=str(DB_PATH), output_dir=None):
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx not installed. Run: pip install python-docx")
        
        self.db_path = db_path
        self.output_dir = output_dir or (DOCS_DIR / "Weather_Policies")
        self.output_dir.mkdir(parents=True, exist_ok=True)
        self.ollama = OllamaService()
        if self.ollama.is_available():
            print(f"🤖 Connected to Local Ollama LLM ({self.ollama.model}) for Weather Policy Synthesis")
    
    def generate_all_policies(self) -> List[str]:
        """Generate weather policy documents from all weather alerts"""
        print("\n" + "="*80)
        print("🌤️  WEATHER POLICY GENERATOR")
        print("="*80)
        
        # Get weather alerts from database
        alerts = self._fetch_weather_alerts()
        print(f"\n📊 Found {len(alerts)} weather alerts in database")
        
        if not alerts:
            print("⚠️  No weather alerts found. Run the main pipeline first.")
            return []
        
        # Group by city and severity
        policies_by_city = self._group_alerts_by_city(alerts)
        
        # Generate policy documents
        generated_files = []
        print(f"\n📝 Generating policy documents...")
        
        for city, city_alerts in policies_by_city.items():
            doc_path = self._create_city_weather_policy(city, city_alerts)
            generated_files.append(doc_path)
            print(f"   ✅ {doc_path.name}")
        
        # Create master weather protocol document
        master_doc = self._create_master_weather_protocol(alerts)
        generated_files.append(master_doc)
        print(f"   ✅ {master_doc.name}")
        
        print(f"\n✅ Generated {len(generated_files)} weather policy documents")
        print(f"📁 Saved to: {self.output_dir}")
        print("="*80 + "\n")
        
        return [str(f) for f in generated_files]
    
    def _fetch_weather_alerts(self) -> List[Dict]:
        """Fetch all weather alerts from database"""
        conn = sqlite3.connect(self.db_path)
        conn.row_factory = sqlite3.Row
        cursor = conn.cursor()
        
        query = """
            SELECT city_name as city, temperature as temp_c, weather_description as description, 
                   wind_speed as wind_ms, rain_1h as rain_mm, visibility_km, 
                   recorded_at as timestamp, 0 as latitude, 0 as longitude
            FROM weather_readings
            WHERE temperature > 40 OR wind_speed > 15 OR rain_1h > 20 OR visibility_km < 1
            ORDER BY city_name, recorded_at DESC
        """
        
        cursor.execute(query)
        alerts = [dict(row) for row in cursor.fetchall()]
        conn.close()
        
        return alerts
    
    def _group_alerts_by_city(self, alerts: List[Dict]) -> Dict[str, List[Dict]]:
        """Group alerts by city"""
        by_city = {}
        for alert in alerts:
            city = alert['city']
            if city not in by_city:
                by_city[city] = []
            by_city[city].append(alert)
        return by_city
    
    def _create_city_weather_policy(self, city: str, alerts: List[Dict]) -> Path:
        """Create a structured, high-density weather policy document for a specific city"""
        doc = Document()
        
        # Telemetry Aggregation
        heat_alerts = [a for a in alerts if a['temp_c'] > 40]
        wind_alerts = [a for a in alerts if a['wind_ms'] > 15]
        rain_alerts = [a for a in alerts if a['rain_mm'] > 20]
        vis_alerts = [a for a in alerts if a['visibility_km'] < 1]
        
        peak_temp = max([a['temp_c'] for a in alerts]) if alerts else 30.0
        peak_wind = max([a['wind_ms'] for a in alerts]) if alerts else 5.0
        peak_rain = max([a['rain_mm'] for a in alerts]) if alerts else 0.0
        min_vis = min([a['visibility_km'] for a in alerts]) if alerts else 10.0
        
        # Title & Metadata Block
        title = doc.add_heading(f'{city.upper()} SEVERE WEATHER PROTOCOL & FREIGHT ADJUDICATION MATRIX', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph(
            f"Monitored Logistics Hub: {city} Commercial Logistics Corridor\n"
            f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')} | Total Telemetry Events Analyzed: {len(alerts)}\n"
            f"Observed Telemetry Extremes: Peak Temp: {peak_temp:.1f}°C | Peak Wind: {peak_wind:.1f} m/s | Peak Rain: {peak_rain:.1f} mm/h | Min Visibility: {min_vis:.2f} km"
        )
        
        # 1. Extracted Meteorological Insights
        doc.add_heading('1. EXTRACTED METEOROLOGICAL INSIGHTS & EXPOSURE PROFILE', 1)
        
        ai_extracted = False
        if self.ollama and self.ollama.is_available():
            alert_summary = [f"- Temp={a.get('temp_c')}C, Wind={a.get('wind_ms')}m/s, Rain={a.get('rain_mm')}mm/h, Cond={a.get('description', '')}" for a in alerts[:5]]
            prompt = (
                f"Analyze these verified meteorological sensor readings for {city}, India:\n"
                + "\n".join(alert_summary) + "\n\n"
                f"TASK: Extract exactly 3 concise, factual insight bullets for freight linehauls in {city}:\n"
                "• Primary Hazard Vector: (State primary hazard e.g. severe crosswind shear, extreme heatwave, or heavy downpour with peak number)\n"
                "• Secondary Hazard Vector: (State secondary atmospheric stress e.g. humidity, slick roads, or visibility drop)\n"
                "• Critical Risk Window: (State freight impact on vehicle stability, linehaul speed reduction, or reefer cooling load)"
            )
            try:
                ai_analysis = self.ollama.generate(prompt)
                if ai_analysis and "Primary Hazard Vector" in ai_analysis:
                    doc.add_paragraph(ai_analysis.strip())
                    ai_extracted = True
            except Exception:
                pass
        
        if not ai_extracted:
            # Deterministic Fallback Insights
            primary_hazard = f"Extreme Heatwave ({peak_temp:.1f}°C)" if peak_temp > 40 else (f"Gale Force Crosswinds ({peak_wind:.1f} m/s)" if peak_wind > 15 else f"Heavy Rain ({peak_rain:.1f} mm/h)")
            doc.add_paragraph(
                f"• Primary Hazard Vector: {primary_hazard} recorded in {city} logistics cluster.\n"
                f"• Secondary Hazard Vector: Ambient atmospheric stress causing transit velocity degradation and cargo vulnerability.\n"
                f"• Critical Risk Window: Elevated highway exposure, linehaul velocity reduction (-25% to -40%), and thermal/moisture packaging stress."
            )
        
        # 2. Logistics Corridor Risk Matrix Table
        doc.add_heading('2. LOGISTICS CORRIDOR & CHOKE POINT RISK MATRIX', 1)
        
        corridors_by_city = {
            "Hyderabad": [
                ("Hyderabad Outer Ring Road (ORR)", f"🔴 HIGH ({peak_wind:.1f} m/s wind)" if peak_wind > 15 else "🟢 LOW", "High-cube trailer rollover hazard", "Enforce 40 km/h speed cap; divert high-cube trailers (>3.0m) to ground arterials."),
                ("Shamshabad Air Cargo Terminal", "🟡 MODERATE", "Crosswind shear & ramp transfer dwell", "Inject +4.0h buffer on air-to-road freight transfers."),
                ("NH-65 (Hyderabad-Vijayawada)", "🟡 MODERATE", "Pavement slickness / heavy linehaul traffic", "Mandatory secondary 80-gauge poly-stretch pallet wrapping."),
                ("Medchal Industrial Logistics Park", "🟢 LOW", "Warehouse staging operations", "Standard operating procedures apply.")
            ],
            "Mumbai": [
                ("NH-48 (Mumbai-Pune Expressway)", f"🔴 SEVERE ({peak_temp:.1f}°C heat / rain)" if peak_temp > 38 else "🟡 MODERATE", "Ghat incline thermal load & engine overheating", "Enforce mandatory reefer pre-cooling to 4°C prior to departure."),
                ("JNPT Port Container Terminal", "🟡 MODERATE", "CFS terminal gate congestion & trailer dwell", "Authorize intermodal drayage bypass if gate queue >4 hours."),
                ("Bhiwandi Central Warehouse Cluster", "🔴 HIGH", "Corrugated carton moisture saturation", "Apply mandatory secondary pallet shrink wrapping; inspect dock seals."),
                ("Western Express Highway Corridor", "🟡 MODERATE", "Urban transit curfew hours (after 08:00)", "Route linehauls via Eastern Freeway or night dispatch windows.")
            ],
            "Bangalore": [
                ("NH-44 (Hosur Road / Electronic City)", "🔴 HIGH", "Arterial toll plaza bottlenecks & gridlock", "Inject +6.0h safety buffer for linehauls connecting to Tamil Nadu."),
                ("Nelamangala Transshipment Yard", "🟡 MODERATE", "Multi-stop LTL consolidation queue", "Enforce FTL direct routing for critical veterinary diets."),
                ("Devanahalli Airport Logistics Belt", "🟢 LOW", "Air freight cold-chain staging", "Standard pre-conditioned reefer transfer protocol.")
            ],
            "Chennai": [
                ("NH-16 (Chennai-Kolkata Corridor)", f"🔴 SEVERE ({peak_rain:.1f} mm/h)" if peak_rain > 15 else "🟡 MODERATE", "Coastal monsoon flooding & road submergence", "Divert linehauls to elevated bypass; enforce SAP QA Hold Stock 'S'."),
                ("Sriperumbudur Industrial Hub", "🟡 MODERATE", "Heavy commercial container traffic", "Schedule dispatch during off-peak windows (22:00-06:00)."),
                ("Chennai Port Trust Terminal Gates", "🟡 MODERATE", "Coastal high humidity & container salt spray", "Mandatory desiccants inside shipping containers.")
            ],
            "Pune": [
                ("NH-48 (Pune-Bangalore Highway)", "🟡 MODERATE", "Chakan belt linehaul congestion", "Add +4.0h dynamic buffer to predicted delivery time."),
                ("Talegaon Logistics Park", "🟢 LOW", "Warehouse consolidation staging", "Standard operating procedures apply.")
            ]
        }
        
        city_corridors = corridors_by_city.get(city, [
            (f"{city} Primary National Highway Corridor", "🟡 MODERATE", "Severe weather exposure & linehaul velocity drop", "Enforce dynamic transit safety buffers and driver weather advisories."),
            (f"{city} Regional Transshipment Hub", "🟢 LOW", "Freight cross-docking operations", "Standard operating procedures apply.")
        ])
        
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Logistics Corridor / Choke Point'
        hdr_cells[1].text = 'Hazard Level'
        hdr_cells[2].text = 'Vulnerability Profile'
        hdr_cells[3].text = 'Mandatory Fleet Directive'
        
        for corridor, haz, vuln, directive in city_corridors:
            row_cells = table.add_row().cells
            row_cells[0].text = corridor
            row_cells[1].text = haz
            row_cells[2].text = vuln
            row_cells[3].text = directive
            
        doc.add_paragraph()  # Spacing
        
        # 3. Binding Operational & QA Directives (Discrete Rules)
        doc.add_heading('3. BINDING OPERATIONAL & QUALITY ASSURANCE (QA) DIRECTIVES', 1)
        
        city_code = city[:3].upper()
        rules = [
            f"[RULE-W-{city_code}-01] TRAILER EQUIPMENT & FLEET RESTRICTION:\n"
            f"  Wind speed >= 15.0 m/s or heavy rain >= 20.0 mm/h mandates immediate suspension of high-cube curtain-sided trailers (>3.0m height). Dispatch must substitute low-profile rigid trucks.",
            
            f"[RULE-W-{city_code}-02] FORCE MAJEURE & SLA ADJUDICATION (CLAUSE 4.2):\n"
            f"  Weather conditions exceeding official alert thresholds qualify as IMD-recognized Severe Meteorological Events. Carrier delay penalties ($500/day -> $0.00) are waived upon submission of verified telematics logs.",
            
            f"[RULE-W-{city_code}-03] COLD-CHAIN & MOISTURE INTEGRITY PROTOCOL:\n"
            f"  Ambient temperature >40.0°C exceeding 4.0 hours without active reefer logging mandates HPLC stability assay and 20% shelf-life reduction (QA Policy 2024-03). Carton moisture >12% mandates 100% rejection.",
            
            f"[RULE-W-{city_code}-04] DYNAMIC ETA BUFFER & REDELIVERY FEE WAIVER:\n"
            f"  Automatically inject a +4.0h to +8.0h dynamic safety buffer into predicted arrival time (PDD). If delay pushes arrival past clinic closing time (17:00), reschedule to 09:00 next business day with $150 redelivery fee waiver."
        ]
        
        for r in rules:
            doc.add_paragraph(r)
            
        # 4. Copilot Deterministic Action Checklist
        doc.add_heading('4. COPILOT DETERMINISTIC ACTION CHECKLIST', 1)
        checklist = [
            f"[x] Step 1: Query live weather telemetry in {city} against alert thresholds (Temp >40°C, Wind >15m/s, Rain >20mm/h).",
            f"[x] Step 2: Cross-reference active SAP linehauls (sap_vttk) entering {city} and check carrier equipment profile.",
            f"[x] Step 3: Apply [RULE-W-{city_code}-02] Force Majeure relief to waive SLA penalties if alert conditions verified.",
            f"[x] Step 4: Dispatch automated Actionable Adaptive Card to Regional Logistics Director for approvals exceeding $500."
        ]
        for item in checklist:
            doc.add_paragraph(item)
            
        # Save Document
        filename = f"{city}_Weather_Protocol.docx"
        doc_path = self.output_dir / filename
        doc.save(str(doc_path))
        return doc_path
    
    def _create_master_weather_protocol(self, all_alerts: List[Dict]) -> Path:
        """Create master weather protocol document covering all cities"""
        doc = Document()
        
        title = doc.add_heading('Master Severe Weather Protocol', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_heading('Purpose', 1)
        doc.add_paragraph(
            "This master protocol provides cross-regional guidance for the O2C Delivery "
            "Risk Copilot when severe weather threatens delivery operations across India."
        )
        
        doc.add_heading('National Weather Impact Summary', 1)
        
        # Get unique cities
        cities = set(a['city'] for a in all_alerts)
        doc.add_paragraph(f"Cities Monitored: {len(cities)}")
        doc.add_paragraph(f"Total Weather Events Analyzed: {len(all_alerts)}")
        
        # Force Majeure Guidelines
        doc.add_heading('Force Majeure Eligibility Criteria', 1)
        doc.add_paragraph(
            "Weather events qualify for Force Majeure protection under "
            "Carrier Master Vendor Agreements when:"
        )
        
        criteria = [
            "Temperature exceeds 42°C (Level 5 Heat Alert)",
            "Wind speed exceeds 20 m/s (Level 4+ Storm)",
            "Rainfall exceeds 50mm/hr (Level 5 Monsoon Alert)",
            "Visibility drops below 0.5km (Level 5 Fog Alert)",
            "Government-issued transport advisory is active",
        ]
        
        for criterion in criteria:
            doc.add_paragraph(criterion, style='List Bullet')
        
        doc.add_paragraph(
            "\nNote: Carrier must provide proof of weather impact at time of delay. "
            "Copilot validates claims against weather API timestamps."
        )
        
        # Regional Routing
        doc.add_heading('Regional Rerouting Matrix', 1)
        doc.add_paragraph(
            "When primary routes are weather-impacted, use these alternatives:"
        )
        
        routes = [
            "Mumbai → Delhi: If Mumbai flooded, route via Pune → Ahmedabad → Delhi",
            "Chennai → Bangalore: If heavy rain, use NH-44 southern corridor",
            "Kolkata → Eastern deliveries: Cyclone season (May-Oct) requires 48hr buffer",
        ]
        
        for route in routes:
            doc.add_paragraph(route, style='List Bullet')
        
        # Save
        doc_path = self.output_dir / "Master_Weather_Protocol.docx"
        doc.save(str(doc_path))
        
        return doc_path


if __name__ == "__main__":
    generator = WeatherPolicyGenerator()
    files = generator.generate_all_policies()
    print(f"\n📄 Generated files:")
    for f in files:
        print(f"   - {f}")
